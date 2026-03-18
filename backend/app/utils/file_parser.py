from __future__ import annotations

from functools import lru_cache
from io import BytesIO
import os
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
from typing import Any

import pdfplumber
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas


IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".webp"}


@lru_cache(maxsize=1)
def _get_paddle_ocr_engine():
    # PaddleX performs online model-source checks by default; disable for offline containers.
    os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
    from paddleocr import PaddleOCR

    return PaddleOCR(use_angle_cls=True, lang="ch")


@lru_cache(maxsize=1)
def _get_rapidocr_engine():
    from rapidocr_onnxruntime import RapidOCR

    return RapidOCR()


def _extract_pdf_text(content: bytes) -> str:
    text_blocks: list[str] = []
    saw_text_layer = False
    with pdfplumber.open(BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = _extract_page_text_with_layout(page)
            if page_text.strip():
                text_blocks.append(page_text.strip())
                saw_text_layer = True

    if saw_text_layer and text_blocks:
        return "\n\n".join(text_blocks)

    # Fallback for scanned PDFs: render pages as images then OCR.
    try:
        import pypdfium2 as pdfium

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as temp_pdf:
            temp_pdf.write(content)
            temp_pdf.flush()
            doc = pdfium.PdfDocument(temp_pdf.name)
            ocr_blocks: list[str] = []
            for page_index in range(len(doc)):
                page = doc[page_index]
                pil_image = page.render(scale=2).to_pil()
                image_buf = BytesIO()
                pil_image.save(image_buf, format="PNG")
                page_text = _extract_image_text(image_buf.getvalue())
                if page_text.strip():
                    ocr_blocks.append(page_text.strip())
            return "\n\n".join(ocr_blocks)
    except Exception:
        return ""


def _extract_page_text_with_layout(page: pdfplumber.page.Page) -> str:
    words = page.extract_words(
        use_text_flow=False,
        keep_blank_chars=False,
        x_tolerance=2,
        y_tolerance=3,
    )
    if not words:
        return ""

    lines: list[dict[str, Any]] = []
    for word in words:
        text = str(word.get("text") or "").strip()
        if not text:
            continue
        y = float(word.get("top", 0.0))
        x = float(word.get("x0", 0.0))
        bucket = None
        for line in lines:
            if abs(line["y"] - y) <= 3.0:
                bucket = line
                break
        if bucket is None:
            bucket = {"y": y, "items": []}
            lines.append(bucket)
        bucket["items"].append((x, text))

    lines.sort(key=lambda line: line["y"])

    rendered_lines: list[str] = []
    for line in lines:
        line_items = sorted(line["items"], key=lambda item: item[0])
        rendered = " ".join(text for _, text in line_items).strip()
        if rendered:
            rendered_lines.append(rendered)

    if not rendered_lines:
        return ""

    rendered_lines = _strip_noise_lines(rendered_lines)
    return "\n".join(rendered_lines)


def _strip_noise_lines(lines: list[str]) -> list[str]:
    def is_noise(line: str) -> bool:
        normalized = line.strip()
        if not normalized:
            return True
        if re.fullmatch(r"[-_\s\d/|.]+", normalized):
            return True
        if re.fullmatch(r"第\s*\d+\s*页", normalized):
            return True
        if re.search(r"版权所有|仅供|机密|保密", normalized):
            return True
        return False

    # Remove header/footer candidates first.
    trimmed = lines[:]
    if trimmed and is_noise(trimmed[0]):
        trimmed = trimmed[1:]
    if trimmed and is_noise(trimmed[-1]):
        trimmed = trimmed[:-1]

    return [line for line in trimmed if not is_noise(line)]


def _extract_docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    blocks: list[str] = []
    for p in document.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        # Preserve rough heading structure for downstream LLM analysis.
        if p.style and str(p.style.name).lower().startswith("heading"):
            blocks.append(f"## {text}")
        else:
            blocks.append(text)
    return "\n".join(blocks)


def _extract_doc_text(content: bytes) -> str:
    # Legacy .doc parsing is unreliable without external office tooling.
    return content.decode("utf-8", errors="ignore").strip()


@lru_cache(maxsize=1)
def _register_cjk_font() -> str:
    font_name = "STSong-Light"
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    except Exception:
        font_name = "Helvetica"
    return font_name


def _text_to_pdf(content: str) -> bytes:
    text = content.strip() or "文档内容为空或暂不支持解析该格式。"
    buffer = BytesIO()
    page_width, page_height = A4
    pdf = canvas.Canvas(buffer, pagesize=A4)
    font_name = _register_cjk_font()
    font_size = 11
    line_height = 18
    left = 40
    top = page_height - 48
    bottom = 40
    max_chars = 46

    pdf.setFont(font_name, font_size)
    y = top
    lines: list[str] = []
    for block in text.splitlines() or [text]:
        block = block.strip()
        if not block:
            lines.append("")
            continue
        while len(block) > max_chars:
            lines.append(block[:max_chars])
            block = block[max_chars:]
        lines.append(block)

    for line in lines:
        if y < bottom:
            pdf.showPage()
            pdf.setFont(font_name, font_size)
            y = top
        pdf.drawString(left, y, line)
        y -= line_height

    pdf.save()
    return buffer.getvalue()


def _try_office_to_pdf(file_name: str, content: bytes) -> bytes | None:
    if shutil.which("soffice") is None:
        return None

    suffix = Path(file_name).suffix.lower()
    if suffix not in {".doc", ".docx"}:
        return None

    with tempfile.TemporaryDirectory() as temp_dir:
        input_path = Path(temp_dir) / f"input{suffix}"
        input_path.write_bytes(content)

        try:
            subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    temp_dir,
                    str(input_path),
                ],
                check=True,
                capture_output=True,
                text=True,
            )
        except Exception:
            return None

        output_path = input_path.with_suffix(".pdf")
        if output_path.exists():
            return output_path.read_bytes()
        return None


def convert_resume_to_preview_pdf(file_name: str, content: bytes) -> bytes:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".pdf":
        return content

    converted = _try_office_to_pdf(file_name, content)
    if converted is not None:
        return converted

    if suffix == ".docx":
        return _text_to_pdf(_extract_docx_text(content))

    if suffix == ".doc":
        return _text_to_pdf(_extract_doc_text(content))

    return _text_to_pdf(content.decode("utf-8", errors="ignore"))


def _extract_texts_from_paddle_result(result: Any) -> list[str]:
    if result is None:
        return []

    if isinstance(result, dict):
        texts = result.get("rec_texts") or result.get("texts") or []
        return [str(text).strip() for text in texts if str(text).strip()]

    if hasattr(result, "res"):
        return _extract_texts_from_paddle_result(result.res)

    if hasattr(result, "json"):
        return _extract_texts_from_paddle_result(result.json)

    if isinstance(result, (list, tuple)):
        text_lines: list[str] = []
        for item in result:
            if isinstance(item, (list, tuple)) and len(item) >= 2 and isinstance(item[1], (list, tuple)):
                candidate = item[1][0] if item[1] else None
                if candidate:
                    text_lines.append(str(candidate).strip())
                continue
            text_lines.extend(_extract_texts_from_paddle_result(item))
        return [line for line in text_lines if line]

    return []


def _extract_text_with_paddle(image_path: str) -> str:
    ocr = _get_paddle_ocr_engine()
    if hasattr(ocr, "predict"):
        result = ocr.predict(image_path)
    else:
        result = ocr.ocr(image_path, cls=True)
    segments = _extract_ocr_segments(result)
    if segments:
        return _compose_segments_by_layout(segments)
    return "\n".join(_extract_texts_from_paddle_result(result))


def _extract_text_with_rapidocr(image_path: str) -> str:
    ocr = _get_rapidocr_engine()
    result, _ = ocr(image_path)
    if not result:
        return ""

    segments = _extract_ocr_segments(result)
    if segments:
        return _compose_segments_by_layout(segments)

    text_lines: list[str] = []
    for item in result:
        if len(item) >= 2 and item[1]:
            text_lines.append(str(item[1]).strip())
    return "\n".join([line for line in text_lines if line])


def _extract_ocr_segments(result: Any) -> list[tuple[float, float, str]]:
    segments: list[tuple[float, float, str]] = []

    def walk(node: Any) -> None:
        if node is None:
            return

        if isinstance(node, dict):
            texts = node.get("rec_texts") or node.get("texts") or []
            boxes = node.get("rec_boxes") or node.get("boxes") or []
            if isinstance(texts, list) and isinstance(boxes, list) and len(texts) == len(boxes):
                for txt, box in zip(texts, boxes):
                    x, y = _top_left_from_box(box)
                    text = str(txt).strip()
                    if text:
                        segments.append((y, x, text))
            for child in node.values():
                walk(child)
            return

        if hasattr(node, "res"):
            walk(node.res)
            return
        if hasattr(node, "json"):
            walk(node.json)
            return

        if isinstance(node, (list, tuple)):
            if len(node) >= 2:
                first, second = node[0], node[1]
                if _looks_like_box(first):
                    text = ""
                    if isinstance(second, str):
                        text = second
                    elif isinstance(second, (list, tuple)) and second:
                        text = str(second[0])
                    x, y = _top_left_from_box(first)
                    text = text.strip()
                    if text:
                        segments.append((y, x, text))
                        return
            for item in node:
                walk(item)

    walk(result)
    return segments


def _looks_like_box(value: Any) -> bool:
    if not isinstance(value, (list, tuple)) or not value:
        return False
    if isinstance(value[0], (list, tuple)):
        return True
    if len(value) == 4 and all(isinstance(v, (int, float)) for v in value):
        return True
    return False


def _top_left_from_box(box: Any) -> tuple[float, float]:
    if isinstance(box, (list, tuple)) and box and isinstance(box[0], (list, tuple)):
        xs = [float(point[0]) for point in box if isinstance(point, (list, tuple)) and len(point) >= 2]
        ys = [float(point[1]) for point in box if isinstance(point, (list, tuple)) and len(point) >= 2]
        if xs and ys:
            return min(xs), min(ys)
    if isinstance(box, (list, tuple)) and len(box) >= 2:
        return float(box[0]), float(box[1])
    return 0.0, 0.0


def _compose_segments_by_layout(segments: list[tuple[float, float, str]]) -> str:
    if not segments:
        return ""

    segments.sort(key=lambda item: (item[0], item[1]))
    lines: list[tuple[float, list[tuple[float, str]]]] = []
    for y, x, text in segments:
        bucket = None
        for ly, items in lines:
            if abs(ly - y) <= 8:
                bucket = items
                break
        if bucket is None:
            bucket = []
            lines.append((y, bucket))
        bucket.append((x, text))

    rendered: list[str] = []
    for _, items in lines:
        items.sort(key=lambda item: item[0])
        line = " ".join(text for _, text in items).strip()
        if line:
            rendered.append(line)
    return "\n".join(rendered)


def _extract_image_text(content: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".png", delete=True) as temp_file:
        temp_file.write(content)
        temp_file.flush()
        errors: list[Exception] = []

        try:
            text = _extract_text_with_paddle(temp_file.name)
            if text.strip():
                return text
        except Exception as exc:
            errors.append(exc)

        try:
            text = _extract_text_with_rapidocr(temp_file.name)
            if text.strip():
                return text
        except Exception as exc:
            errors.append(exc)

    if errors:
        raise RuntimeError("image OCR failed") from errors[-1]
    return ""


def _normalize_resume_text(text: str) -> str:
    text = text.replace("\u3000", " ").replace("\xa0", " ")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\t\f\v]+", " ", text)
    text = re.sub(r"[ ]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Normalize common OCR punctuation artifacts.
    text = text.replace("，", ",").replace("；", ";").replace("：", ":")
    text = re.sub(r"[�□◆◇※]+", "", text)

    # Preserve heading/list cues while cleaning noisy symbols.
    cleaned_lines: list[str] = []
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            cleaned_lines.append("")
            continue
        line = re.sub(r"^[•·●○]+\s*", "- ", line)
        cleaned_lines.append(line)

    normalized = "\n".join(cleaned_lines)
    return normalized.strip()


def _ensure_pdf_bytes(file_name: str, content: bytes) -> tuple[bytes, str]:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".pdf":
        return content, "pdf-native"

    if suffix not in {".doc", ".docx"}:
        return b"", ""

    converted = _try_office_to_pdf(file_name, content)
    if converted is not None:
        return converted, "office-pdf"

    text = _extract_docx_text(content) if suffix == ".docx" else _extract_doc_text(content)
    return _text_to_pdf(text), "text-pdf"


def parse_resume_file(file_name: str, content: bytes) -> dict:
    suffix = Path(file_name).suffix.lower()
    mode = "text"

    if suffix == ".pdf":
        text = _extract_pdf_text(content)
        mode = "pdf"
    elif suffix == ".docx":
        pdf_bytes, converted_by = _ensure_pdf_bytes(file_name, content)
        text = _extract_pdf_text(pdf_bytes)
        mode = f"docx->{converted_by or 'pdf'}"
    elif suffix == ".doc":
        pdf_bytes, converted_by = _ensure_pdf_bytes(file_name, content)
        text = _extract_pdf_text(pdf_bytes)
        mode = f"doc->{converted_by or 'pdf'}"
    elif suffix in IMAGE_SUFFIXES:
        text = _extract_image_text(content)
        mode = "image-ocr"
    else:
        text = content.decode("utf-8", errors="ignore")
        mode = "plain-text"

    cleaned = _normalize_resume_text(text)
    return {
        "file_name": file_name,
        "size": len(content),
        "mode": mode,
        "suffix": suffix,
        "text": cleaned,
    }
