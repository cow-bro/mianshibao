from __future__ import annotations

import json
import tempfile
import time
from pathlib import Path

import requests
from docx import Document
from PIL import Image, ImageDraw
from sqlalchemy import select

from app.core.database import SessionLocal
from app.core.security import get_password_hash
from app.models.user import User, UserRole

API_BASE = "http://backend:8000/api/v1"


def login_and_get_auth_header() -> dict[str, str]:
    code, resp = call(
        "POST",
        f"{API_BASE}/auth/login",
        json={"username": "demo", "password": "demo123"},
    )
    if code != 200:
        raise RuntimeError(f"login failed: status={code}, body={resp.text}")

    payload = resp.json()
    access_token = payload.get("data", {}).get("access_token")
    if not access_token:
        raise RuntimeError(f"missing access token in login response: {payload}")

    return {"Authorization": f"Bearer {access_token}"}


async def ensure_demo_user() -> int:
    async with SessionLocal() as db:
        result = await db.execute(select(User).where(User.username == "demo"))
        user = result.scalar_one_or_none()
        if user is None:
            user = User(
                username="demo",
                hashed_password=get_password_hash("demo123"),
                role=UserRole.USER,
                is_active=True,
            )
            db.add(user)
            await db.commit()
            await db.refresh(user)
        else:
            user.hashed_password = get_password_hash("demo123")
            user.is_active = True
            if user.role is None:
                user.role = UserRole.USER
            await db.commit()
        return user.id


def create_sample_files(work_dir: Path) -> dict[str, Path]:
    docx_path = work_dir / "stage4_resume.docx"
    png_path = work_dir / "stage4_resume.png"
    pdf_path = work_dir / "stage4_resume.pdf"

    doc = Document()
    doc.add_heading("Demo Resume", level=1)
    doc.add_paragraph("Name: Demo User")
    doc.add_paragraph("Email: demo@example.com")
    doc.add_paragraph("Phone: 13812345678")
    doc.add_paragraph("Experience: Built async FastAPI service with PostgreSQL and MinIO.")
    doc.save(docx_path)

    img = Image.new("RGB", (900, 450), "white")
    draw = ImageDraw.Draw(img)
    draw.text((30, 50), "Demo Resume", fill="black")
    draw.text((30, 120), "Email: demo@example.com", fill="black")
    draw.text((30, 190), "Phone: 13812345678", fill="black")
    draw.text((30, 260), "Project: resume parser and scorer", fill="black")
    img.save(png_path)
    img.save(pdf_path, "PDF")

    return {"docx": docx_path, "png": png_path, "pdf": pdf_path}


def call(method: str, url: str, timeout: int = 120, retries: int = 30, **kwargs):
    last_error = None
    for _ in range(retries):
        try:
            resp = requests.request(method=method, url=url, timeout=timeout, **kwargs)
            return resp.status_code, resp
        except (requests.exceptions.ConnectionError, requests.exceptions.Timeout) as exc:
            last_error = exc
            time.sleep(1)
    raise RuntimeError(f"request failed after retries: {url}") from last_error


def main() -> None:
    import asyncio

    asyncio.run(ensure_demo_user())
    auth_headers = login_and_get_auth_header()

    report: dict[str, dict] = {}

    # JWT protection check: upload without token should fail.
    with tempfile.TemporaryDirectory() as tmp:
        files = create_sample_files(Path(tmp))

        with files["docx"].open("rb") as f:
            code, resp = call("POST", f"{API_BASE}/resumes/upload", files={"file": f})
        report["jwt_protection"] = {
            "status_code": code,
            "passed": code == 401,
            "message": "upload without token should return 401",
        }

        uploaded: dict[str, int] = {}

        for kind, path in files.items():
            with path.open("rb") as f:
                code, resp = call(
                    "POST",
                    f"{API_BASE}/resumes/upload",
                    headers=auth_headers,
                    files={"file": (path.name, f)},
                )
            data = {}
            if resp.headers.get("content-type", "").startswith("application/json"):
                data = resp.json()
            passed = code == 200 and data.get("code") == 0 and bool(data.get("data", {}).get("resume_id"))
            report[f"upload_{kind}"] = {"status_code": code, "passed": passed, "data": data}
            if passed:
                uploaded[kind] = int(data["data"]["resume_id"])

        for kind, resume_id in uploaded.items():
            code, resp = call(
                "POST",
                f"{API_BASE}/resumes/{resume_id}/parse",
                headers=auth_headers,
                timeout=600,
                retries=3,
            )
            data = resp.json() if resp.headers.get("content-type", "").startswith("application/json") else {}
            parsed_data = data.get("data") if isinstance(data, dict) else {}
            if not isinstance(parsed_data, dict):
                parsed_data = {}
            has_minimal = bool(
                parsed_data.get("personal_info", {}).get("email")
                or parsed_data.get("personal_info", {}).get("phone")
            )
            report[f"parse_{kind}"] = {
                "status_code": code,
                "passed": code == 200 and data.get("code") == 0 and has_minimal,
                "data": data,
            }

        # Full chain on docx if available.
        if "docx" in uploaded:
            docx_id = uploaded["docx"]

            code, resp = call("POST", f"{API_BASE}/resumes/{docx_id}/score", headers=auth_headers)
            data = resp.json() if resp.headers.get("content-type", "").startswith("application/json") else {}
            score_data = data.get("data", {}) if isinstance(data, dict) else {}
            dims = score_data.get("dimension_scores", {})
            report["score_docx"] = {
                "status_code": code,
                "passed": code == 200 and data.get("code") == 0 and bool(score_data.get("overall_score")) and len(dims) >= 4,
                "data": data,
            }

            code, resp = call("POST", f"{API_BASE}/resumes/{docx_id}/optimize", headers=auth_headers)
            data = resp.json() if resp.headers.get("content-type", "").startswith("application/json") else {}
            opt_data = data.get("data", {}) if isinstance(data, dict) else {}
            report["optimize_docx"] = {
                "status_code": code,
                "passed": code == 200 and data.get("code") == 0 and bool(opt_data.get("optimized_content")),
                "data": data,
            }

            code, resp = call(
                "GET",
                f"{API_BASE}/resumes/{docx_id}/download-optimized",
                headers=auth_headers,
            )
            content_disposition = resp.headers.get("content-disposition", "")
            report["download_optimized_docx"] = {
                "status_code": code,
                "passed": code == 200 and "attachment" in content_disposition.lower() and len(resp.content) > 0,
                "content_disposition": content_disposition,
                "bytes": len(resp.content),
            }

    failed = [k for k, v in report.items() if not v.get("passed")]
    summary = {
        "total": len(report),
        "passed": len(report) - len(failed),
        "failed": len(failed),
        "failed_items": failed,
    }

    print(json.dumps({"summary": summary, "report": report}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
